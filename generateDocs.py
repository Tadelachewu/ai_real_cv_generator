from datetime import datetime
import os
import traceback
from typing import Dict
from venv import logger

from docx import Document
from docx.shared import Inches
from jinja2 import Environment, FileSystemLoader
from urllib.request import pathname2url  # ✅ Proper path conversion for URI

from ai import enhance_with_ai
import logging  # ✅ Correct logger module

logger = logging.getLogger(__name__)  # ✅ Proper logger setup

# from bot import TEMP_DIR, TEMPLATES_DIR, enhance_with_ai
# Initialize base directory
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_DIR = os.path.join(BASE_DIR, 'temp')
TEMPLATES_DIR = os.path.join(BASE_DIR, 'templates')

# Ensure directories exist
os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

def generate_pdf(data: Dict) -> str:
    """Generate PDF CV with robust error handling and photo support"""
    try:
        enhanced_data = data.copy()

        # ✅ Handle photo path correctly for WeasyPrint
        photo_path = enhanced_data.get("photo_path")
        if photo_path and os.path.exists(photo_path):
            abs_path = os.path.abspath(photo_path)
            # Use pathname2url to ensure special characters and slashes are properly encoded
            enhanced_data['photo_url'] = f"file://{pathname2url(abs_path)}"
        else:
            enhanced_data['photo_url'] = None
            logger.info("No valid photo included for PDF")

        # ✅ Set up template environment
        env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))
        template_name = enhanced_data.get('template', 'professional')
        template_file = f"{template_name}_cv.html"

        try:
            template = env.get_template(template_file)
        except Exception as e:
            logger.warning(f"Template {template_file} not found. Falling back to default. Reason: {e}")
            template = env.get_template("professional_cv.html")

        # ✅ Add optional fields if missing
        enhanced_data.update({
            'current_date': datetime.now().strftime("%B %Y"),
            'linkedin': enhanced_data.get('linkedin', ''),
            'portfolio': enhanced_data.get('portfolio', ''),
            'skills': enhanced_data.get('skills', []),
            'languages': enhanced_data.get('languages', []),
            'certifications': enhanced_data.get('certifications', []),
            'projects': enhanced_data.get('projects', []),
            'experience': enhanced_data.get('experience', []),
            'education': enhanced_data.get('education', [])
        })

        # ✅ Render the HTML from Jinja2 template
        html = template.render(enhanced_data)

        # ✅ Create a safe filename for the PDF
        safe_name = "".join(c if c.isalnum() or c in (' ', '_') else '_' for c in enhanced_data.get('name', 'cv'))
        filename = os.path.join(TEMP_DIR, f"{safe_name}_CV.pdf")

        # ✅ Generate PDF using WeasyPrint (import lazily to avoid import-time failures)
        try:
            from weasyprint import HTML, CSS
        except Exception as e:
            logger.error("WeasyPrint import failed when generating PDF: %s", e)
            raise RuntimeError(
                "WeasyPrint cannot be used in this environment. Install GTK/cairo/pango/gdk-pixbuf or run in Docker.\n"
                f"Original error: {e}"
            )

        html_obj = HTML(string=html)
        css = CSS(string='@page { size: A4; margin: 1cm; }')
        html_obj.write_pdf(filename, stylesheets=[css])

        logger.info(f"✅ PDF successfully generated at: {filename}")
        return filename

    except Exception as e:
        logger.error(f"❌ PDF generation failed: {type(e).__name__}: {str(e)}")
        logger.error(f"Traceback:\n{traceback.format_exc()}")
        raise RuntimeError(f"Failed to generate PDF: {str(e)}")
    
def generate_docx(data: Dict) -> str:
    """Generate DOCX version of CV with improved reliability"""
    try:
        enhanced_data = data.copy()  # Don't enhance with AI here to keep it fast
        
        # Create document
        doc = Document()
        
        # Add photo if exists
        photo_path = enhanced_data.get("photo_path")
        if photo_path and os.path.exists(photo_path):
            try:
                doc.add_picture(photo_path, width=Inches(1.5))
                doc.add_paragraph()  # Add space after photo
            except Exception as e:
                logger.warning(f"Couldn't add photo: {e}")

        # Header section
        name = enhanced_data.get('name', 'Your Name')
        doc.add_heading(name, level=0)
        
        # Contact information
        contact_info = []
        if enhanced_data.get('email'):
            contact_info.append(enhanced_data['email'])
        if enhanced_data.get('phone'):
            contact_info.append(enhanced_data['phone'])
        
        contact = doc.add_paragraph(' | '.join(contact_info))
        contact.alignment = 1  # Center alignment
        
        # Optional links
        if enhanced_data.get('linkedin'):
            doc.add_paragraph(f"LinkedIn: {enhanced_data['linkedin']}", style='BodyText')
        if enhanced_data.get('portfolio'):
            doc.add_paragraph(f"Portfolio: {enhanced_data['portfolio']}", style='BodyText')

        # Summary section
        if enhanced_data.get('summary'):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(enhanced_data['summary'])

        # Experience section
        if enhanced_data.get('experience'):
            doc.add_heading('Professional Experience', level=1)
            for job in enhanced_data['experience']:
                # Job title and company
                p = doc.add_paragraph()
                p.add_run(f"{job.get('role', 'Position')}").bold = True
                p.add_run(f" at {job.get('company', 'Company')} | {job.get('years', 'Dates')}")
                
                # Job description
                if job.get('description'):
                    doc.add_paragraph(job['description'], style='List Bullet')

        # Education section
        if enhanced_data.get('education'):
            doc.add_heading('Education', level=1)
            for edu in enhanced_data['education']:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', 'Degree')}").bold = True
                p.add_run(f" | {edu.get('institution', 'Institution')} | {edu.get('years', 'Dates')}")

        # Skills section
        if enhanced_data.get('skills'):
            doc.add_heading('Skills', level=1)
            if isinstance(enhanced_data['skills'], dict):
                for category, skills in enhanced_data['skills'].items():
                    doc.add_paragraph(f"{category.capitalize()}:", style='Heading 3')
                    doc.add_paragraph(', '.join(skills))
            else:
                doc.add_paragraph(', '.join(str(skill) for skill in enhanced_data['skills']))
                # Languages section
        if enhanced_data.get('languages'):
            doc.add_heading('Languages', level=1)
            if isinstance(enhanced_data['languages'], list):
                doc.add_paragraph(', '.join(str(lang) for lang in enhanced_data['languages']))
            elif isinstance(enhanced_data['languages'], dict):
                for lang, level in enhanced_data['languages'].items():
                    doc.add_paragraph(f"{lang}: {level}")

        # Projects section
        if enhanced_data.get('projects'):
            doc.add_heading('Projects', level=1)
            for project in enhanced_data['projects']:
                p = doc.add_paragraph()
                p.add_run(f"{project.get('name', 'Project')}").bold = True
                if project.get('technologies'):
                    p.add_run(f" | Technologies: {project['technologies']}")
                if project.get('description'):
                    doc.add_paragraph(project['description'], style='List Bullet')

        # Generate filename and save
        safe_name = "".join(c if c.isalnum() or c in (' ', '_') else '_' for c in name)
        filename = os.path.join(TEMP_DIR, f"{safe_name}_CV.docx")
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        
        doc.save(filename)
        logger.info(f"DOCX successfully generated at: {filename}")
        return filename
        
    except Exception as e:
        logger.error(f"DOCX generation failed: {type(e).__name__}: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise RuntimeError(f"Failed to generate DOCX: {str(e)}")